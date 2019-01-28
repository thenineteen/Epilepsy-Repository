"""
Read word document texts for patients with
Epilepsy, read drugs table, remove sensitive data,
pseudo-anonymise and store new popped text
and original in separate folders

example to call this function:
path_to_doc = 'L:\word_docs\word_docx_folder\FILENAME.docx'
pt_docx, pt_meds_list = epilepsy_docx(path_to_doc, 4, 7, read_tables=True)

Ali Alim-Marvasti (c) Jan 2019
"""
# below only works for docx
# so first convert .doc to .docx via .ps1
import docx


def args_for_loop(*args):
    list = []
    for arg in args:
        list.append(arg)
    return list


def epilepsy_docx(path_to_doc, *paragraphs, read_tables=False):
    """
    prints the text of .docx file.
    paragraphs: optional,
      reads between two paragraph numbers (inclusive).
      if not specified, or too many numbers specified, reads all paragraphs.
      if only one number specified, reads from paragraph 0
        to the number specified.
    read_tables will read the tables as paragraphs, if True.
    Returns pt_docx as list of paragraphs.
    Returns pt_meds_list as list of the table - needs further cleanup.
    """
    document = docx.Document(path_to_doc)
    pt_docx = []  # initialise patient's docx list
    pt_meds_list = []  # initialise patient meds list

    if paragraphs:
        para_list = args_for_loop(*paragraphs)

    if len(paragraphs) == 2:
        for n, p in enumerate(document.paragraphs):
            if n >= para_list[0] and n <= para_list[1]:
                pt_docx.append(p.text)
                print(p.text)

    elif len(paragraphs) == 1:
        for n, p in enumerate(document.paragraphs):
            if n >= 0 and n <= para_list[0]:
                pt_docx.append(p.text)
                print(p.text)

    else:
        for p in document.paragraphs:
            pt_docx.append(p.text)
            print(p.text)

    if read_tables:
        for t in document.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        pt_meds_list.append(paragraph.text)
                        # print(paragraph.text)
        # print (pt_meds_list)

    # this part cleans the pt_meds_list to remove empty characters and spaces
    pt_meds_list_clean_and_lower = [
        itm.lower() for itm in pt_meds_list if itm != '' and itm != '\t']
    pt_meds_list_clean_phenytoin = [
        s.replace('  ', '') for s in pt_meds_list_clean_and_lower]
    pt_meds_list_clean_phenytoin = [med.replace(
        'phenytoin ', 'phenytoin') for med in pt_meds_list_clean_phenytoin]
    pt_meds_list = pt_meds_list_clean_phenytoin  # simpler rename

    return pt_docx, pt_meds_list
